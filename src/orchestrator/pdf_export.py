from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List, Union

# ---- Optional imports (lazy-friendly)
try:
    # Only needed when backend="reportlab"
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
except Exception:
    colors = LETTER = ParagraphStyle = getSampleStyleSheet = inch = SimpleDocTemplate = Paragraph = Spacer = Table = TableStyle = None  # type: ignore


# ------------------------- ReportLab helpers ---------------------------

def _reportlab_styles():
    styles = getSampleStyleSheet()
    TitleStyle = ParagraphStyle("TitleStyle", parent=styles["Heading1"], spaceAfter=12)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=12, spaceAfter=6)
    H3 = ParagraphStyle("H3", parent=styles["Heading3"], spaceBefore=6, spaceAfter=4)
    Body = ParagraphStyle("Body", parent=styles["BodyText"], leading=14, spaceAfter=6)
    Small = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=9, leading=12)
    return TitleStyle, H2, H3, Body, Small


def _kv_table_reportlab(section_i: Dict[str, Any], Body) -> "Table":
    rows: List[List[Any]] = []
    for k, v in section_i.items():
        key = k[:-1] if k.endswith(":") else k
        rows.append([Paragraph(f"<b>{key}</b>", Body), Paragraph(str(v or ""), Body)])
    t = Table(rows, colWidths=[2.2 * inch, 4.8 * inch])
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
            ]
        )
    )
    return t


def _add_clause_block_reportlab(story: List[Any], block: Dict[str, Any], H3, Body) -> None:
    fields = [
        ("Clause Language", "clause_language"),
        ("Clause Summary", "clause_summary"),
        ("Risk Triggers", "risk_triggers"),
        ("Flow-Down Obligations", "flow_down_obligations"),
        ("Redline Recommendations", "redline_recommendations"),
        ("Harmful Language / Conflicts", "harmful_language_conflicts"),
    ]
    for label, key in fields:
        val = block.get(key)
        if val:
            story.append(Paragraph(f"<b>{label}</b>", H3))
            story.append(Paragraph(str(val), Body))


def _section_items_reportlab(story: List[Any], section_name: str, items: Iterable[Dict[str, Any]], H2, H3, Body, Small) -> None:
    story.append(Paragraph(section_name, H2))
    empty = True
    for item in items or []:
        empty = False
        title = item.get("item_title") or f"{section_name} Item"
        story.append(Paragraph(title, H3))
        for block in item.get("clauses") or []:
            _add_clause_block_reportlab(story, block, H3, Body)
        story.append(Spacer(1, 6))
    if empty:
        story.append(Paragraph("No items found.", Small))


def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.drawRightString(LETTER[0] - 0.5 * inch, 0.5 * inch, f"Page {doc.page}")
    canvas.restoreState()


def _export_reportlab(data: Dict[str, Any], output_pdf: Path, title: str) -> Path:
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab is not installed; install it or use backend='docx'")
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    TitleStyle, H2, H3, Body, Small = _reportlab_styles()

    doc = SimpleDocTemplate(
        str(output_pdf),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title=title,
        author="CR2A",
    )

    story: List[Any] = []
    story.append(Paragraph(title, TitleStyle))
    story.append(Spacer(1, 6))

    section_i = data.get("SECTION_I") or {}
    story.append(Paragraph("SECTION I — Contract Overview", H2))
    story.append(_kv_table_reportlab(section_i, Body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("SECTION II — Parties & Dates", H2))
    _section_items_reportlab(story, "Parties & Dates", data.get("SECTION_II") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION III — Scope & Deliverables", H2))
    _section_items_reportlab(story, "Scope & Deliverables", data.get("SECTION_III") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION IV — Payment Terms", H2))
    _section_items_reportlab(story, "Payment Terms", data.get("SECTION_IV") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION V — Risk Allocation", H2))
    _section_items_reportlab(story, "Risk Allocation", data.get("SECTION_V") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION VI — Compliance & Legal", H2))
    _section_items_reportlab(story, "Compliance & Legal", data.get("SECTION_VI") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return output_pdf


# ---------------------------- DOCX backend -----------------------------

def _export_docx(data: Dict[str, Any], template_docx: Path, output_pdf: Path, title: str) -> Path:
    from docx import Document
    try:
        from docx2pdf import convert as docx2pdf_convert
    except Exception:
        docx2pdf_convert = None  # optional

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    if template_docx and template_docx.exists():
        doc = Document(str(template_docx))
    else:
        doc = Document()
        doc.add_heading(title, level=1)

    # SECTION I
    doc.add_heading("SECTION I — Contract Overview", level=2)
    section_i = data.get("SECTION_I") or {}
    table = doc.add_table(rows=0, cols=2)
    for k, v in section_i.items():
        row = table.add_row().cells
        row[0].text = (k[:-1] if k.endswith(":") else k)
        row[1].text = str(v or "")

    def write_items(hdr: str, items: Iterable[Dict[str, Any]]):
        doc.add_heading(hdr, level=2)
        wrote = False
        for item in items or []:
            wrote = True
            doc.add_heading(str(item.get("item_title") or "Item"), level=3)
            for block in item.get("clauses") or []:
                for label, key in [
                    ("Clause Language", "clause_language"),
                    ("Clause Summary", "clause_summary"),
                    ("Risk Triggers", "risk_triggers"),
                    ("Flow-Down Obligations", "flow_down_obligations"),
                    ("Redline Recommendations", "redline_recommendations"),
                    ("Harmful Language / Conflicts", "harmful_language_conflicts"),
                ]:
                    val = block.get(key)
                    if val:
                        p = doc.add_paragraph()
                        run = p.add_run(f"{label}: ")
                        run.bold = True
                        doc.add_paragraph(str(val))
        if not wrote:
            doc.add_paragraph("No items found.")

    write_items("SECTION II — Parties & Dates", data.get("SECTION_II"))
    write_items("SECTION III — Scope & Deliverables", data.get("SECTION_III"))
    write_items("SECTION IV — Payment Terms", data.get("SECTION_IV"))
    write_items("SECTION V — Risk Allocation", data.get("SECTION_V"))
    write_items("SECTION VI — Compliance & Legal", data.get("SECTION_VI"))

    out_docx = output_pdf.with_suffix(".docx")
    doc.save(str(out_docx))

    # Convert to PDF if possible; otherwise leave the .docx
    try:
        from docx2pdf import convert as docx2pdf_convert2
        if docx2pdf_convert2:
            docx2pdf_convert2(str(out_docx), str(output_pdf))
            return output_pdf
    except Exception:
        pass
    return out_docx


# ------------------------------ Public API -----------------------------

def export_pdf_from_filled_json(
    data: Dict[str, Any],
    output_pdf: Union[str, Path],
    *,
    backend: str = "reportlab",  # default to ReportLab to avoid OS-level conversion
    template_docx: Union[str, Path, None] = None,
    title: str = "Contract Risk & Compliance Analysis",
) -> Path:
    """
    Render a CR2A filled JSON into a PDF (or DOCX if conversion is unavailable).
    backend="docx" uses python-docx (+ docx2pdf if available); backend="reportlab" uses ReportLab.
    """
    output_pdf = Path(output_pdf).expanduser().resolve()
    backend = (backend or "reportlab").lower()
    if backend == "docx":
        tpl = Path(template_docx) if template_docx else Path("")
        return _export_docx(data, tpl, output_pdf, title)
    return _export_reportlab(data, output_pdf, title)
