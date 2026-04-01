"""
Contract Uploader Module

Handles file validation and text extraction for contract documents.
Supports PDF and DOCX file formats with optional OCR for image-based PDFs.
"""

import logging
import os
import sys
from pathlib import Path
from typing import Tuple, Dict, Optional
try:
    from pypdf import PdfReader
except ImportError:
    # Fallback to PyPDF2 if pypdf is not available
    from PyPDF2 import PdfReader
from docx import Document

# Excel support
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# Optional OCR support
try:
    import pytesseract
    from PIL import Image
    import pdf2image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False


logger = logging.getLogger(__name__)


class ContractUploader:
    """
    Handles contract file validation and text extraction.
    
    This class provides functionality to:
    - Validate file formats (PDF/DOCX only)
    - Extract metadata (filename, size, page count)
    - Extract text content from PDF and DOCX files
    """
    
    # Supported file extensions
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.xlsx'}
    
    def __init__(self, max_file_size: Optional[int] = None, enable_ocr: bool = True, tesseract_path: Optional[str] = None):
        """
        Initialize the Contract Uploader.
        
        Args:
            max_file_size: Maximum file size in bytes (default: 250 MB for backward compatibility)
            enable_ocr: Enable OCR for image-based PDFs (default: True if Tesseract available)
            tesseract_path: Path to Tesseract executable (optional, auto-detected if not provided)
        """
        self.MAX_FILE_SIZE = max_file_size if max_file_size is not None else 250 * 1024 * 1024
        self.enable_ocr = enable_ocr and TESSERACT_AVAILABLE
        
        # Configure Tesseract if available
        if self.enable_ocr and tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            logger.info("Tesseract OCR enabled with custom path: %s", tesseract_path)
        elif self.enable_ocr:
            # Auto-detect bundled Tesseract for frozen (PyInstaller) builds
            bundled_tesseract = self._find_bundled_tesseract()
            if bundled_tesseract:
                pytesseract.pytesseract.tesseract_cmd = bundled_tesseract
                logger.info("Tesseract OCR enabled with bundled path: %s", bundled_tesseract)
            else:
                logger.info("Tesseract OCR enabled (using system PATH)")
        else:
            if not TESSERACT_AVAILABLE:
                logger.info("Tesseract OCR not available (missing dependencies: pytesseract, pillow, pdf2image)")
            else:
                logger.info("Tesseract OCR disabled by configuration")
        
        logger.debug("ContractUploader initialized with max file size: %d bytes (%.2f MB)", 
                     self.MAX_FILE_SIZE, self.MAX_FILE_SIZE / (1024 * 1024))
    
    @staticmethod
    def _find_bundled_tesseract() -> Optional[str]:
        """Find bundled Tesseract executable in frozen or development builds."""
        import sys
        candidates = []
        if getattr(sys, 'frozen', False):
            base_path = sys._MEIPASS if hasattr(sys, '_MEIPASS') else os.path.dirname(sys.executable)
            candidates.extend([
                os.path.join(base_path, 'tesseract', 'tesseract.exe'),
                os.path.join(os.path.dirname(sys.executable), '_internal', 'tesseract', 'tesseract.exe'),
            ])
        # Also check build_tools/ relative to project root (development mode)
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        candidates.append(os.path.join(project_root, 'build_tools', 'tesseract', 'tesseract.exe'))
        for path in candidates:
            if os.path.exists(path):
                # Also set TESSDATA_PREFIX so tesseract finds language files
                # TESSDATA_PREFIX must point to the tessdata/ directory itself
                tessdata_dir = os.path.join(os.path.dirname(path), 'tessdata')
                if os.path.exists(tessdata_dir):
                    os.environ['TESSDATA_PREFIX'] = tessdata_dir
                return path
        return None

    def validate_format(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate file format (PDF/DOCX).
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            Tuple of (is_valid, error_message)
            - is_valid: True if file format is supported, False otherwise
            - error_message: Empty string if valid, error description if invalid
        """
        logger.debug("Validating file format: %s", file_path)
        
        try:
            # Convert to Path object for better handling
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                error_msg = f"File not found: {file_path}"
                logger.warning(error_msg)
                return False, error_msg
            
            # Check if it's a file (not a directory)
            if not path.is_file():
                error_msg = f"Path is not a file: {file_path}"
                logger.warning(error_msg)
                return False, error_msg
            
            # Check file extension
            file_extension = path.suffix.lower()
            if file_extension not in self.SUPPORTED_FORMATS:
                error_msg = (
                    f"Unsupported file format: {file_extension}\n"
                    f"Supported formats: {', '.join(sorted(self.SUPPORTED_FORMATS))}"
                )
                logger.warning(error_msg)
                return False, error_msg
            
            # Check file size
            file_size = path.stat().st_size
            if file_size == 0:
                error_msg = "File is empty (0 bytes)"
                logger.warning(error_msg)
                return False, error_msg
            
            if file_size > self.MAX_FILE_SIZE:
                error_msg = (
                    f"File size ({file_size / (1024*1024):.1f} MB) exceeds "
                    f"maximum allowed size ({self.MAX_FILE_SIZE / (1024*1024):.0f} MB)"
                )
                logger.warning(error_msg)
                return False, error_msg
            
            # Check if file is readable
            try:
                with open(path, 'rb') as f:
                    # Try to read first byte to verify readability
                    f.read(1)
            except PermissionError:
                error_msg = f"Permission denied: Cannot read file {file_path}"
                logger.warning(error_msg)
                return False, error_msg
            except Exception as e:
                error_msg = f"Cannot read file: {str(e)}"
                logger.warning(error_msg)
                return False, error_msg
            
            logger.info("File format validation passed: %s", file_path)
            return True, ""
            
        except Exception as e:
            error_msg = f"Validation error: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg
    
    def get_file_info(self, file_path: str) -> Dict[str, any]:
        """
        Get file metadata (name, size, page count).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing file metadata:
            - filename: Name of the file
            - file_size_bytes: Size in bytes
            - file_size_mb: Size in megabytes (formatted)
            - page_count: Number of pages (for PDF) or None
            - file_type: File extension
        """
        logger.debug("Extracting file info: %s", file_path)
        
        try:
            path = Path(file_path)
            
            # Basic file information
            file_info = {
                'filename': path.name,
                'file_size_bytes': path.stat().st_size,
                'file_size_mb': f"{path.stat().st_size / (1024*1024):.2f} MB",
                'file_type': path.suffix.lower(),
                'page_count': None
            }
            
            # Try to get page count for PDF files
            if path.suffix.lower() == '.pdf':
                try:
                    with open(path, 'rb') as f:
                        pdf_reader = PdfReader(f)
                        file_info['page_count'] = len(pdf_reader.pages)
                        logger.debug("PDF page count: %d", file_info['page_count'])
                except Exception as e:
                    logger.warning("Could not extract page count from PDF: %s", e)
                    file_info['page_count'] = None
            
            # Try to get page count for DOCX files (approximate based on paragraphs)
            elif path.suffix.lower() == '.docx':
                try:
                    doc = Document(path)
                    # Rough estimate: ~30 paragraphs per page
                    paragraph_count = len(doc.paragraphs)
                    file_info['page_count'] = max(1, paragraph_count // 30)
                    logger.debug("DOCX estimated page count: %d", file_info['page_count'])
                except Exception as e:
                    logger.warning("Could not extract page count from DOCX: %s", e)
                    file_info['page_count'] = None
            
            logger.info("File info extracted successfully: %s", file_info['filename'])
            return file_info
            
        except Exception as e:
            logger.error("Error extracting file info: %s", e, exc_info=True)
            # Return minimal info on error
            return {
                'filename': os.path.basename(file_path),
                'file_size_bytes': 0,
                'file_size_mb': "Unknown",
                'file_type': Path(file_path).suffix.lower(),
                'page_count': None
            }

    def extract_text(self, file_path: str, progress_callback=None) -> str:
        """
        Extract text content from PDF or DOCX file.

        Args:
            file_path: Path to the contract file
            progress_callback: Optional callable(status, percent) for page-level progress

        Returns:
            Extracted text content as a string

        Raises:
            ValueError: If file format is not supported
            Exception: If text extraction fails
        """
        logger.info("Extracting text from file: %s", file_path)

        # Validate file format first
        is_valid, error_msg = self.validate_format(file_path)
        if not is_valid:
            logger.error("File validation failed: %s", error_msg)
            raise ValueError(f"Cannot extract text: {error_msg}")

        path = Path(file_path)
        file_extension = path.suffix.lower()

        try:
            if file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path, progress_callback=progress_callback)
            elif file_extension == '.docx':
                return self._extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_text_from_txt(file_path)
            elif file_extension == '.xlsx':
                return self._extract_text_from_xlsx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error("Text extraction failed for %s: %s", file_path, e, exc_info=True)
            raise
    
    def _extract_text_from_pdf(self, file_path: str, progress_callback=None) -> str:
        """
        Extract text from PDF file using PyPDF2.

        Args:
            file_path: Path to PDF file
            progress_callback: Optional callable(status, percent) for page-level progress

        Returns:
            Extracted text content

        Raises:
            Exception: If PDF reading fails
        """
        logger.debug("Extracting text from PDF: %s", file_path)
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.warning("PDF is encrypted, attempting to decrypt with empty password")
                    try:
                        # Try to decrypt with empty password
                        pdf_reader.decrypt('')
                    except Exception as decrypt_error:
                        error_msg = (
                            "PDF is encrypted and cannot be read. "
                            "Please provide an unencrypted version of the document."
                        )
                        logger.error(error_msg)
                        raise ValueError(error_msg) from decrypt_error
                
                # Extract text from all pages
                text_parts = []
                page_count = len(pdf_reader.pages)
                
                logger.debug("Extracting text from %d pages", page_count)
                
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    if progress_callback:
                        pct = int(100 * page_num / page_count)
                        progress_callback(f"Reading page {page_num}/{page_count}...", pct)
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"\n--- Page {page_num} ---\n")
                            text_parts.append(page_text)
                    except Exception as page_error:
                        logger.warning("Failed to extract text from page %d: %s",
                                     page_num, page_error)
                        # Continue with other pages
                        continue
                
                extracted_text = '\n'.join(text_parts)
                
                if not extracted_text.strip():
                    # No text extracted - try OCR if available
                    fname = file_path.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
                    if self.enable_ocr:
                        logger.info("No text extracted from PDF, attempting OCR...")
                        try:
                            extracted_text = self._extract_text_with_ocr(file_path, progress_callback=progress_callback)
                            logger.info("Successfully extracted %d characters using OCR", len(extracted_text))
                            return extracted_text
                        except Exception as ocr_error:
                            logger.error("OCR extraction failed: %s", ocr_error)
                            error_msg = (
                                f"'{fname}' appears to be a scanned/image PDF.\n"
                                "Text extraction failed and OCR could not process it.\n\n"
                                "This is common for stamped plans and drawings.\n"
                                "Try loading the text-based bid documents instead\n"
                                "(ITB, specifications, addenda).\n\n"
                                f"OCR detail: {str(ocr_error)[:200]}"
                            )
                            logger.error(error_msg)
                            raise ValueError(error_msg) from ocr_error
                    else:
                        error_msg = (
                            f"'{fname}' appears to be a scanned/image PDF.\n"
                            "No text could be extracted.\n\n"
                            "This is common for stamped plans and drawings.\n"
                            "Try loading the text-based bid documents instead\n"
                            "(ITB, specifications, addenda)."
                        )
                        logger.error(error_msg)
                        raise ValueError(error_msg)
                
                logger.info("Successfully extracted %d characters from PDF", 
                          len(extracted_text))
                return extracted_text
                
        except ValueError:
            # Re-raise ValueError (encrypted PDF, no text extracted)
            raise
        except Exception as e:
            error_msg = f"Failed to read PDF file: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise Exception(error_msg) from e
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If DOCX reading fails
        """
        logger.debug("Extracting text from DOCX: %s", file_path)
        
        try:
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            extracted_text = '\n\n'.join(text_parts)
            
            if not extracted_text.strip():
                error_msg = (
                    "No text could be extracted from the DOCX file. "
                    "The document may be empty or corrupted."
                )
                logger.error(error_msg)
                raise ValueError(error_msg)
            
            logger.info("Successfully extracted %d characters from DOCX", 
                      len(extracted_text))
            return extracted_text
            
        except ValueError:
            # Re-raise ValueError (no text extracted)
            raise
        except Exception as e:
            error_msg = f"Failed to read DOCX file: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise Exception(error_msg) from e
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If TXT reading fails
        """
        logger.debug("Extracting text from TXT: %s", file_path)
        
        try:
            # Try UTF-8 first
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()
            except UnicodeDecodeError:
                # Fallback to latin-1 if UTF-8 fails
                logger.warning("UTF-8 decode failed, trying latin-1")
                with open(file_path, 'r', encoding='latin-1') as f:
                    extracted_text = f.read()
            
            if not extracted_text.strip():
                error_msg = "The text file is empty."
                logger.error(error_msg)
                raise ValueError(error_msg)
            
            logger.info("Successfully extracted %d characters from TXT", 
                      len(extracted_text))
            return extracted_text
            
        except ValueError:
            # Re-raise ValueError (empty file)
            raise
        except Exception as e:
            error_msg = f"Failed to read TXT file: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise Exception(error_msg) from e

    def _extract_text_from_xlsx(self, file_path: str) -> str:
        """Extract text from Excel (.xlsx) file."""
        logger.debug("Extracting text from XLSX: %s", file_path)

        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl is required to read Excel files. Install with: pip install openpyxl")

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            all_text = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                all_text.append(f"=== Sheet: {sheet_name} ===")

                for row in ws.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    cells = [str(cell).strip() for cell in row if cell is not None]
                    if cells:
                        all_text.append(" | ".join(cells))

            wb.close()
            extracted_text = "\n".join(all_text)

            if not extracted_text.strip():
                raise ValueError("The Excel file contains no data.")

            logger.info("Successfully extracted %d characters from XLSX (%d sheets)",
                       len(extracted_text), len(wb.sheetnames))
            return extracted_text

        except ValueError:
            raise
        except Exception as e:
            error_msg = f"Failed to read Excel file: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise Exception(error_msg) from e

    def _extract_text_with_ocr(self, file_path: str, progress_callback=None) -> str:
        """
        Extract text from image-based PDF using Tesseract OCR.

        Args:
            file_path: Path to PDF file
            progress_callback: Optional callable(status, percent) for page-level progress

        Returns:
            Extracted text content from OCR

        Raises:
            Exception: If OCR extraction fails
            RuntimeError: If Tesseract is not available
        """
        if not self.enable_ocr:
            raise RuntimeError("OCR is not enabled or Tesseract is not available")

        logger.info("Starting OCR extraction for: %s", file_path)

        try:
            # Detect bundled poppler path for frozen or development builds
            poppler_path = None
            candidates = []
            if getattr(sys, 'frozen', False):
                base_path = getattr(sys, '_MEIPASS', os.path.dirname(sys.executable))
                candidates.extend([
                    os.path.join(base_path, 'poppler', 'bin'),
                    os.path.join(os.path.dirname(sys.executable), '_internal', 'poppler', 'bin'),
                ])
            # Also check build_tools/ relative to project root (development mode)
            project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            candidates.append(os.path.join(project_root, 'build_tools', 'poppler', 'bin'))
            for candidate in candidates:
                if os.path.exists(candidate):
                    poppler_path = candidate
                    logger.debug("Using bundled poppler at: %s", poppler_path)
                    break

            # Convert PDF pages to images
            logger.debug("Converting PDF to images...")
            if poppler_path:
                images = pdf2image.convert_from_path(file_path, poppler_path=poppler_path)
            else:
                images = pdf2image.convert_from_path(file_path)
            logger.info("Converted PDF to %d images", len(images))
            
            if progress_callback:
                progress_callback("Converting PDF to images for OCR...", 5)

            # Extract text from each page image
            text_parts = []
            total_pages = len(images)
            for page_num, image in enumerate(images, start=1):
                logger.debug("Running OCR on page %d/%d", page_num, total_pages)
                if progress_callback:
                    pct = int(100 * page_num / total_pages)
                    progress_callback(f"OCR processing page {page_num}/{total_pages}...", pct)
                
                try:
                    # Run Tesseract OCR on the image
                    page_text = pytesseract.image_to_string(image, lang='eng')
                    
                    if page_text.strip():
                        text_parts.append(f"\n--- Page {page_num} (OCR) ---\n")
                        text_parts.append(page_text)
                        logger.debug("Extracted %d characters from page %d", 
                                   len(page_text), page_num)
                    else:
                        logger.warning("No text extracted from page %d", page_num)
                        
                except Exception as page_error:
                    logger.warning("OCR failed for page %d: %s", page_num, page_error)
                    # Continue with other pages
                    continue
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise ValueError("OCR did not extract any text from the document")
            
            logger.info("OCR extraction complete: %d characters extracted from %d pages", 
                       len(extracted_text), len(images))
            return extracted_text
            
        except Exception as e:
            error_msg = f"OCR extraction failed: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise Exception(error_msg) from e


def find_text_in_pdf(pdf_path: str, search_text: str, page_hint: Optional[int] = None) -> Tuple[Optional[int], Optional[float]]:
    """
    Find the exact PDF coordinates of a text string using PyMuPDF.

    Args:
        pdf_path: Path to the PDF file.
        search_text: Text to search for (typically a section header).
        page_hint: Optional 1-based page number to search first.

    Returns:
        (page_number, y_coordinate) — 1-based page number and top y-position
        in PDF points, or (None, None) if not found.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.debug("PyMuPDF not available for coordinate lookup")
        return (None, None)

    if not search_text or not pdf_path or not os.path.isfile(pdf_path):
        return (None, None)

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        logger.debug("Could not open PDF for coordinate search: %s", e)
        return (None, None)

    try:
        # Clean the search text — strip numbering artifacts, limit length
        clean = search_text.strip()
        if len(clean) > 150:
            clean = clean[:150]

        # Build a list of pages to search: page_hint first (if given), then all others
        pages_to_search = list(range(len(doc)))
        if page_hint is not None and 1 <= page_hint <= len(doc):
            hint_idx = page_hint - 1
            pages_to_search.remove(hint_idx)
            pages_to_search.insert(0, hint_idx)

        # Try the full search text first, then progressively shorter prefixes
        search_variants = [clean]
        # Add a shorter prefix if the header is long (more likely to match)
        if len(clean) > 40:
            search_variants.append(clean[:60])
        if len(clean) > 20:
            search_variants.append(clean[:30])

        for variant in search_variants:
            for page_idx in pages_to_search:
                page = doc[page_idx]
                rects = page.search_for(variant, quads=False)
                if rects:
                    # Return the first match: 1-based page and top y-coordinate
                    rect = rects[0]
                    page_num = page_idx + 1
                    y_top = rect.y0
                    logger.debug("Found '%s' on page %d at y=%.1f", variant[:40], page_num, y_top)
                    doc.close()
                    return (page_num, y_top)

        doc.close()
        logger.debug("Text not found in PDF: '%s'", clean[:40])
        return (None, None)

    except Exception as e:
        logger.debug("Error searching PDF for text coordinates: %s", e)
        try:
            doc.close()
        except Exception:
            pass
        return (None, None)


def page_from_char_position(contract_text: str, char_pos: int) -> Optional[int]:
    """
    Derive the PDF page number from a character position in extracted text.

    Scans for ``--- Page N ---`` markers inserted during extraction and returns
    the page number that contains the given character position.

    Args:
        contract_text: Full extracted text with page markers.
        char_pos: Character offset into the text.

    Returns:
        1-based page number, or None if markers are not present.
    """
    import re
    current_page = None
    for m in re.finditer(r'--- Page (\d+)(?:\s*\(OCR\))? ---', contract_text):
        marker_pos = m.start()
        if marker_pos > char_pos:
            break
        current_page = int(m.group(1))
    return current_page
