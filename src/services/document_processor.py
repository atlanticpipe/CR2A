"""
Document Processing Service
Handles extraction of text from various document formats (PDF, DOCX, TXT).
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document text extraction from multiple formats."""

    SUPPORTED_FORMATS = {"pdf", "docx", "txt"}

    def extract_text(self, filepath: str) -> str:
        """
        Extract text from a document file.

        Args:
            filepath: Path to the document file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported
            IOError: If file cannot be read
        """
        path = Path(filepath)

        if not path.exists():
            raise IOError(f"File not found: {filepath}")

        file_ext = path.suffix.lower().lstrip(".")

        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        logger.info(f"Extracting text from {file_ext.upper()}: {path.name}")

        if file_ext == "pdf":
            return self._extract_from_pdf(filepath)
        elif file_ext == "docx":
            return self._extract_from_docx(filepath)
        elif file_ext == "txt":
            return self._extract_from_txt(filepath)

        raise ValueError(f"Unsupported format: {file_ext}")

    @staticmethod
    def _extract_from_pdf(filepath: str) -> str:
        """
        Extract text from PDF using pdfplumber.

        Args:
            filepath: Path to PDF file

        Returns:
            Extracted text
        """
        if pdfplumber is None:
            raise ImportError("pdfplumber not installed")

        text_parts = []

        try:
            with pdfplumber.open(filepath) as pdf:
                logger.info(f"Processing PDF with {len(pdf.pages)} pages")

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        else:
                            logger.warning(f"Page {page_num} produced no text")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        continue

            combined_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(combined_text)} characters from PDF")
            return combined_text

        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise IOError(f"Failed to extract text from PDF: {e}")

    @staticmethod
    def _extract_from_docx(filepath: str) -> str:
        """
        Extract text from DOCX using python-docx.

        Args:
            filepath: Path to DOCX file

        Returns:
            Extracted text
        """
        if DocxDocument is None:
            raise ImportError("python-docx not installed")

        try:
            doc = DocxDocument(filepath)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            combined_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(combined_text)} characters from DOCX")
            return combined_text

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise IOError(f"Failed to extract text from DOCX: {e}")

    @staticmethod
    def _extract_from_txt(filepath: str) -> str:
        """
        Extract text from plain text file.

        Args:
            filepath: Path to TXT file

        Returns:
            File contents
        """
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()

            logger.info(f"Extracted {len(text)} characters from TXT")
            return text

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(filepath, "r", encoding="latin-1") as f:
                    text = f.read()

                logger.info(f"Extracted {len(text)} characters from TXT (latin-1)")
                return text
            except Exception as e:
                logger.error(f"Error processing TXT: {e}")
                raise IOError(f"Failed to extract text from TXT: {e}")

        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            raise IOError(f"Failed to extract text from TXT: {e}")

    @staticmethod
    def get_file_info(filepath: str) -> dict:
        """
        Get information about a document file.

        Args:
            filepath: Path to the file

        Returns:
            Dictionary with file information
        """
        path = Path(filepath)
        file_ext = path.suffix.lower().lstrip(".")

        return {
            "filename": path.name,
            "filepath": str(path),
            "format": file_ext,
            "size_bytes": path.stat().st_size if path.exists() else 0,
            "size_mb": round(path.stat().st_size / (1024 * 1024), 2)
            if path.exists()
            else 0,
        }
