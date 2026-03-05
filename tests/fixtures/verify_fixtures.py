"""
Script to verify that test fixtures are working correctly.
"""

from PyPDF2 import PdfReader
from docx import Document
import os
import json
from pathlib import Path

def test_valid_pdfs():
    """Test that valid PDF files can be read."""
    print("Testing valid PDF files...")
    pdf_files = [
        'contract_1page.pdf',
        'contract_10pages.pdf',
        'contract_25pages.pdf',
        'contract_50pages.pdf'
    ]
    
    for filename in pdf_files:
        filepath = f'tests/fixtures/{filename}'
        try:
            reader = PdfReader(filepath)
            pages = len(reader.pages)
            text = reader.pages[0].extract_text()
            print(f"  ✓ {filename}: {pages} pages, {len(text)} chars on first page")
        except Exception as e:
            print(f"  ✗ {filename}: ERROR - {e}")

def test_valid_docx():
    """Test that valid DOCX files can be read."""
    print("\nTesting valid DOCX files...")
    docx_files = [
        'contract_1page.docx',
        'contract_10pages.docx',
        'contract_25pages.docx',
        'contract_50pages.docx'
    ]
    
    for filename in docx_files:
        filepath = f'tests/fixtures/{filename}'
        try:
            doc = Document(filepath)
            paragraphs = len(doc.paragraphs)
            text = '\n'.join([p.text for p in doc.paragraphs])
            print(f"  ✓ {filename}: {paragraphs} paragraphs, {len(text)} total chars")
        except Exception as e:
            print(f"  ✗ {filename}: ERROR - {e}")

def test_malformed_files():
    """Test that malformed files raise appropriate errors."""
    print("\nTesting malformed files...")
    
    # Test fake PDF
    try:
        reader = PdfReader('tests/fixtures/malformed_fake.pdf')
        print(f"  ✗ malformed_fake.pdf: Should have raised an error!")
    except Exception as e:
        print(f"  ✓ malformed_fake.pdf: Correctly raised {type(e).__name__}")
    
    # Test fake DOCX
    try:
        doc = Document('tests/fixtures/malformed_fake.docx')
        print(f"  ✗ malformed_fake.docx: Should have raised an error!")
    except Exception as e:
        print(f"  ✓ malformed_fake.docx: Correctly raised {type(e).__name__}")
    
    # Test empty PDF
    try:
        reader = PdfReader('tests/fixtures/empty.pdf')
        print(f"  ✗ empty.pdf: Should have raised an error!")
    except Exception as e:
        print(f"  ✓ empty.pdf: Correctly raised {type(e).__name__}")
    
    # Test corrupted PDF
    try:
        reader = PdfReader('tests/fixtures/corrupted.pdf')
        print(f"  ✗ corrupted.pdf: Should have raised an error!")
    except Exception as e:
        print(f"  ✓ corrupted.pdf: Correctly raised {type(e).__name__}")

def test_file_sizes():
    """Test that files have reasonable sizes."""
    print("\nTesting file sizes...")
    
    files_to_check = [
        ('contract_1page.pdf', 5, 20),
        ('contract_10pages.pdf', 10, 30),
        ('contract_25pages.pdf', 30, 60),
        ('contract_50pages.pdf', 60, 120),
        ('contract_1page.docx', 30, 50),
        ('contract_10pages.docx', 35, 50),
        ('contract_25pages.docx', 35, 50),
        ('contract_50pages.docx', 35, 50),
    ]
    
    for filename, min_kb, max_kb in files_to_check:
        filepath = f'tests/fixtures/{filename}'
        size_kb = os.path.getsize(filepath) / 1024
        if min_kb <= size_kb <= max_kb:
            print(f"  ✓ {filename}: {size_kb:.2f} KB (within {min_kb}-{max_kb} KB range)")
        else:
            print(f"  ✗ {filename}: {size_kb:.2f} KB (expected {min_kb}-{max_kb} KB)")

def test_api_response_fixtures():
    """Test that API response fixtures are valid JSON and have correct structure."""
    print("\nTesting API response fixtures...")
    
    fixtures_dir = Path('tests/fixtures')
    
    # Success response fixtures - should be valid JSON with expected structure
    success_fixtures = [
        'api_response_success_full.json',
        'api_response_success_minimal.json',
        'api_response_empty_arrays.json'
    ]
    
    for filename in success_fixtures:
        filepath = fixtures_dir / filename
        try:
            with open(filepath, 'r') as f:
                data = json.load(f)
            
            # Check for expected top-level keys (some may be missing in partial responses)
            has_metadata = 'contract_metadata' in data
            has_clauses = 'clauses' in data
            has_risks = 'risks' in data
            has_compliance = 'compliance_issues' in data
            has_redlining = 'redlining_suggestions' in data
            
            status = "✓" if all([has_clauses, has_risks, has_compliance, has_redlining]) else "⚠"
            print(f"  {status} {filename}: Valid JSON, metadata={has_metadata}, clauses={has_clauses}, risks={has_risks}")
        except json.JSONDecodeError as e:
            print(f"  ✗ {filename}: Invalid JSON - {e}")
        except Exception as e:
            print(f"  ✗ {filename}: ERROR - {e}")
    
    # Partial response fixtures - valid JSON but missing some fields
    partial_fixtures = [
        'api_response_partial_missing_risks.json',
        'api_response_partial_missing_clauses.json',
        'api_response_partial_missing_metadata.json'
    ]
    
    for filename in partial_fixtures:
        filepath = fixtures_dir / filename
        try:
            with open(filepath, 'r') as f:
                data = json.load(f)
            
            missing_fields = []
            if 'contract_metadata' not in data:
                missing_fields.append('metadata')
            if 'clauses' not in data:
                missing_fields.append('clauses')
            if 'risks' not in data:
                missing_fields.append('risks')
            
            print(f"  ✓ {filename}: Valid JSON, intentionally missing: {', '.join(missing_fields) if missing_fields else 'none'}")
        except json.JSONDecodeError as e:
            print(f"  ✗ {filename}: Invalid JSON - {e}")
        except Exception as e:
            print(f"  ✗ {filename}: ERROR - {e}")
    
    # Error response fixtures - should be valid JSON with error structure
    error_fixtures = [
        'api_error_rate_limit.json',
        'api_error_invalid_key.json',
        'api_error_network_timeout.json',
        'api_error_server_error.json',
        'api_error_invalid_model.json',
        'api_error_context_length.json'
    ]
    
    for filename in error_fixtures:
        filepath = fixtures_dir / filename
        try:
            with open(filepath, 'r') as f:
                data = json.load(f)
            
            has_error = 'error' in data
            if has_error:
                error = data['error']
                has_message = 'message' in error
                has_type = 'type' in error
                has_code = 'code' in error
                has_status = 'status' in error
                
                if all([has_message, has_type, has_code, has_status]):
                    print(f"  ✓ {filename}: Valid error response (status={error['status']})")
                else:
                    print(f"  ⚠ {filename}: Missing error fields")
            else:
                print(f"  ✗ {filename}: Missing 'error' key")
        except json.JSONDecodeError as e:
            print(f"  ✗ {filename}: Invalid JSON - {e}")
        except Exception as e:
            print(f"  ✗ {filename}: ERROR - {e}")
    
    # Malformed fixtures - should fail to parse or have wrong types
    malformed_fixtures = [
        ('api_response_invalid_json.txt', 'should not be valid JSON'),
        ('api_response_malformed_json.json', 'valid JSON but wrong types'),
        ('api_response_incomplete_json.json', 'incomplete/truncated JSON')
    ]
    
    for filename, description in malformed_fixtures:
        filepath = fixtures_dir / filename
        try:
            with open(filepath, 'r') as f:
                data = json.load(f)
            
            if filename == 'api_response_malformed_json.json':
                print(f"  ✓ {filename}: Valid JSON but {description}")
            else:
                print(f"  ⚠ {filename}: Expected to fail but parsed successfully")
        except json.JSONDecodeError:
            print(f"  ✓ {filename}: Correctly fails to parse ({description})")
        except Exception as e:
            print(f"  ✓ {filename}: Raises error as expected - {type(e).__name__}")

def main():
    """Run all verification tests."""
    print("=" * 60)
    print("Verifying Test Fixtures")
    print("=" * 60)
    
    test_valid_pdfs()
    test_valid_docx()
    test_malformed_files()
    test_file_sizes()
    test_api_response_fixtures()
    
    print("\n" + "=" * 60)
    print("Verification Complete!")
    print("=" * 60)

if __name__ == '__main__':
    main()
