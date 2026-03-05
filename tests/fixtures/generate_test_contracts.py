"""
Script to generate sample contract files for testing.

This script creates:
- PDF contracts of various sizes (1, 10, 25, 50 pages)
- DOCX contracts of various sizes
- Malformed files for error testing
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from docx import Document
from docx.shared import Pt, Inches
import os

# Sample contract text with realistic clauses
CONTRACT_TEMPLATE = """
SERVICE AGREEMENT

This Service Agreement ("Agreement") is entered into as of {date}, by and between:

CLIENT: ABC Corporation, a Delaware corporation with its principal place of business at 123 Main Street, 
New York, NY 10001 ("Client")

SERVICE PROVIDER: XYZ Services LLC, a Delaware limited liability company with its principal place of 
business at 456 Oak Avenue, San Francisco, CA 94102 ("Service Provider")

WHEREAS, Client desires to engage Service Provider to provide certain services; and
WHEREAS, Service Provider desires to provide such services to Client;

NOW, THEREFORE, in consideration of the mutual covenants and agreements hereinafter set forth and for 
other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, 
the parties agree as follows:

1. SERVICES

1.1 Scope of Services. Service Provider shall provide the following services to Client (the "Services"):
    (a) Software development and maintenance services
    (b) Technical consulting and advisory services
    (c) System integration and deployment services
    (d) Training and documentation services

1.2 Performance Standards. Service Provider shall perform the Services in a professional and workmanlike 
manner in accordance with industry standards and best practices.

2. TERM AND TERMINATION

2.1 Term. This Agreement shall commence on the Effective Date and shall continue for a period of 
twelve (12) months (the "Initial Term"), unless earlier terminated as provided herein.

2.2 Renewal. This Agreement shall automatically renew for successive one (1) year periods (each a 
"Renewal Term") unless either party provides written notice of non-renewal at least sixty (60) days 
prior to the end of the then-current term.

2.3 Termination for Convenience. Either party may terminate this Agreement for any reason upon ninety 
(90) days prior written notice to the other party.

2.4 Termination for Cause. Either party may terminate this Agreement immediately upon written notice 
if the other party materially breaches this Agreement and fails to cure such breach within thirty (30) 
days after receiving written notice thereof.

3. COMPENSATION AND PAYMENT

3.1 Fees. Client shall pay Service Provider the fees set forth in Exhibit A attached hereto (the "Fees").

3.2 Payment Terms. All Fees shall be due and payable within thirty (30) days of the date of invoice. 
Late payments shall accrue interest at the rate of 1.5% per month or the maximum rate permitted by law, 
whichever is less.

3.3 Expenses. Client shall reimburse Service Provider for all reasonable and documented out-of-pocket 
expenses incurred in connection with the performance of the Services, provided that such expenses are 
pre-approved by Client in writing.

4. INTELLECTUAL PROPERTY

4.1 Ownership. All work product, deliverables, and materials created by Service Provider in the course 
of performing the Services (collectively, "Work Product") shall be the sole and exclusive property of 
Client.

4.2 Assignment. Service Provider hereby assigns to Client all right, title, and interest in and to the 
Work Product, including all intellectual property rights therein.

4.3 Pre-Existing Materials. Service Provider retains all right, title, and interest in and to any 
pre-existing materials, tools, utilities, and methodologies used in connection with the Services.

5. CONFIDENTIALITY

5.1 Confidential Information. Each party acknowledges that it may have access to certain confidential 
and proprietary information of the other party, including but not limited to trade secrets, business 
plans, financial information, customer lists, and technical data (collectively, "Confidential Information").

5.2 Non-Disclosure. Each party agrees to maintain the confidentiality of the other party's Confidential 
Information and not to disclose such information to any third party without the prior written consent 
of the disclosing party.

5.3 Exceptions. The obligations of confidentiality shall not apply to information that: (a) is or 
becomes publicly available through no breach of this Agreement; (b) is rightfully received from a third 
party without breach of any confidentiality obligation; (c) is independently developed without use of 
the Confidential Information; or (d) is required to be disclosed by law or court order.

6. WARRANTIES AND DISCLAIMERS

6.1 Service Provider Warranties. Service Provider warrants that: (a) it has the right and authority to 
enter into this Agreement and perform the Services; (b) the Services will be performed in a professional 
and workmanlike manner; and (c) the Work Product will not infringe any third party intellectual property 
rights.

6.2 Client Warranties. Client warrants that it has the right and authority to enter into this Agreement 
and to provide Service Provider with access to any materials, systems, or information necessary for the 
performance of the Services.

6.3 DISCLAIMER. EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, SERVICE PROVIDER MAKES NO WARRANTIES, 
EXPRESS OR IMPLIED, INCLUDING WITHOUT LIMITATION ANY IMPLIED WARRANTIES OF MERCHANTABILITY, FITNESS FOR 
A PARTICULAR PURPOSE, OR NON-INFRINGEMENT.

7. LIMITATION OF LIABILITY

7.1 Cap on Liability. IN NO EVENT SHALL EITHER PARTY'S TOTAL LIABILITY ARISING OUT OF OR RELATED TO 
THIS AGREEMENT EXCEED THE TOTAL FEES PAID OR PAYABLE TO SERVICE PROVIDER UNDER THIS AGREEMENT IN THE 
TWELVE (12) MONTHS PRECEDING THE EVENT GIVING RISE TO LIABILITY.

7.2 Exclusion of Consequential Damages. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, 
INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING WITHOUT LIMITATION LOST PROFITS, 
LOST REVENUE, OR LOST DATA, EVEN IF SUCH PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.

7.3 Exceptions. The limitations set forth in this Section 7 shall not apply to: (a) either party's 
indemnification obligations; (b) either party's breach of confidentiality obligations; or (c) either 
party's gross negligence or willful misconduct.

8. INDEMNIFICATION

8.1 Service Provider Indemnification. Service Provider shall indemnify, defend, and hold harmless Client 
from and against any and all claims, damages, losses, and expenses (including reasonable attorneys' fees) 
arising out of or related to: (a) Service Provider's breach of this Agreement; (b) Service Provider's 
negligence or willful misconduct; or (c) any claim that the Work Product infringes any third party 
intellectual property rights.

8.2 Client Indemnification. Client shall indemnify, defend, and hold harmless Service Provider from and 
against any and all claims, damages, losses, and expenses (including reasonable attorneys' fees) arising 
out of or related to: (a) Client's breach of this Agreement; (b) Client's negligence or willful misconduct; 
or (c) any claim arising from Client's use of the Work Product.

9. GENERAL PROVISIONS

9.1 Governing Law. This Agreement shall be governed by and construed in accordance with the laws of the 
State of Delaware, without regard to its conflicts of law principles.

9.2 Dispute Resolution. Any dispute arising out of or related to this Agreement shall be resolved through 
binding arbitration in accordance with the Commercial Arbitration Rules of the American Arbitration 
Association. The arbitration shall be conducted in Wilmington, Delaware.

9.3 Entire Agreement. This Agreement constitutes the entire agreement between the parties with respect 
to the subject matter hereof and supersedes all prior and contemporaneous agreements and understandings, 
whether written or oral.

9.4 Amendments. This Agreement may not be amended or modified except by a written instrument signed by 
both parties.

9.5 Waiver. No waiver of any provision of this Agreement shall be deemed or shall constitute a waiver 
of any other provision, nor shall any waiver constitute a continuing waiver.

9.6 Severability. If any provision of this Agreement is held to be invalid or unenforceable, the remaining 
provisions shall continue in full force and effect.

9.7 Assignment. Neither party may assign this Agreement without the prior written consent of the other 
party, except that either party may assign this Agreement to a successor in connection with a merger, 
acquisition, or sale of all or substantially all of its assets.

9.8 Notices. All notices under this Agreement shall be in writing and shall be deemed given when delivered 
personally, sent by confirmed facsimile, sent by confirmed email, or three (3) days after being sent by 
certified mail, return receipt requested, to the addresses set forth above.

9.9 Force Majeure. Neither party shall be liable for any failure or delay in performance due to causes 
beyond its reasonable control, including but not limited to acts of God, war, terrorism, riots, embargoes, 
acts of civil or military authorities, fire, floods, accidents, strikes, or shortages of transportation, 
facilities, fuel, energy, labor, or materials.

9.10 Independent Contractors. The parties are independent contractors and nothing in this Agreement shall 
be construed to create a partnership, joint venture, agency, or employment relationship between the parties.

9.11 Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an 
original and all of which together shall constitute one and the same instrument.

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

CLIENT:                                    SERVICE PROVIDER:
ABC Corporation                            XYZ Services LLC

By: _________________________              By: _________________________
Name: John Smith                           Name: Jane Doe
Title: Chief Executive Officer             Title: Managing Director
Date: _______________________              Date: _______________________
"""

ADDITIONAL_CLAUSE_TEMPLATE = """
{section_num}. ADDITIONAL TERMS AND CONDITIONS

{section_num}.1 Data Protection and Privacy. Service Provider shall comply with all applicable data 
protection and privacy laws, including but not limited to the General Data Protection Regulation (GDPR) 
and the California Consumer Privacy Act (CCPA). Service Provider shall implement appropriate technical 
and organizational measures to protect personal data processed in connection with the Services.

{section_num}.2 Security Requirements. Service Provider shall maintain commercially reasonable security 
measures to protect Client's systems, data, and Confidential Information from unauthorized access, use, 
or disclosure. Service Provider shall promptly notify Client of any security breach or unauthorized 
access to Client's systems or data.

{section_num}.3 Insurance. Service Provider shall maintain, at its own expense, the following insurance 
coverage: (a) commercial general liability insurance with limits of not less than $2,000,000 per 
occurrence; (b) professional liability insurance with limits of not less than $2,000,000 per claim; 
and (c) workers' compensation insurance as required by law.

{section_num}.4 Audit Rights. Client shall have the right, upon reasonable notice and during normal 
business hours, to audit Service Provider's performance of the Services and compliance with this Agreement. 
Service Provider shall cooperate with such audits and provide Client with access to relevant records and 
personnel.

{section_num}.5 Subcontractors. Service Provider may engage subcontractors to perform portions of the 
Services only with Client's prior written consent. Service Provider shall remain fully responsible for 
the performance of any subcontractors and shall ensure that all subcontractors comply with the terms of 
this Agreement.

{section_num}.6 Service Level Agreement. Service Provider shall meet the service levels set forth in 
Exhibit B attached hereto (the "SLA"). If Service Provider fails to meet the SLA, Client shall be entitled 
to the remedies set forth in the SLA, which shall be Client's sole and exclusive remedy for such failure.

{section_num}.7 Change Management. Any changes to the scope of Services must be documented in a written 
change order signed by both parties. The change order shall specify the nature of the change, any impact 
on the Fees, and any impact on the schedule for performance of the Services.

{section_num}.8 Compliance with Laws. Each party shall comply with all applicable federal, state, and 
local laws, regulations, and ordinances in connection with its performance under this Agreement, including 
but not limited to laws relating to employment, discrimination, occupational health and safety, and 
environmental protection.
"""


def generate_pdf_contract(filename, num_pages):
    """Generate a PDF contract with the specified number of pages."""
    doc = SimpleDocTemplate(filename, pagesize=letter,
                          topMargin=1*inch, bottomMargin=1*inch,
                          leftMargin=1*inch, rightMargin=1*inch)
    
    story = []
    styles = getSampleStyleSheet()
    
    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='black',
        spaceAfter=30,
        alignment=1  # Center
    )
    
    # Add main contract text
    body_style = styles['BodyText']
    body_style.fontSize = 10
    body_style.leading = 14
    
    # Add the main contract
    for line in CONTRACT_TEMPLATE.split('\n'):
        if line.strip():
            p = Paragraph(line, body_style)
            story.append(p)
            story.append(Spacer(1, 0.1*inch))
    
    # Add additional sections to reach the desired page count
    section_num = 10
    while len(story) < num_pages * 40:  # Approximate paragraphs per page
        additional_text = ADDITIONAL_CLAUSE_TEMPLATE.format(section_num=section_num)
        for line in additional_text.split('\n'):
            if line.strip():
                p = Paragraph(line, body_style)
                story.append(p)
                story.append(Spacer(1, 0.1*inch))
        section_num += 1
        
        # Add some filler content if needed
        if len(story) < num_pages * 40:
            filler = f"""
{section_num}. SUPPLEMENTARY PROVISIONS

{section_num}.1 This section contains additional terms and conditions that supplement the main agreement. 
The parties acknowledge that these provisions are essential to the proper execution and performance of 
this Agreement. Each party represents and warrants that it has carefully reviewed these provisions and 
understands their implications.

{section_num}.2 The parties further agree that any ambiguity in this Agreement shall be resolved in 
accordance with the principles of contract interpretation applicable in the governing jurisdiction. 
Neither party shall be deemed to be the drafter of this Agreement for purposes of interpreting any 
ambiguous provisions.
"""
            for line in filler.split('\n'):
                if line.strip():
                    p = Paragraph(line, body_style)
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))
            section_num += 1
    
    doc.build(story)
    print(f"Generated PDF: {filename} ({num_pages} pages)")


def generate_docx_contract(filename, num_pages):
    """Generate a DOCX contract with the specified number of pages."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('SERVICE AGREEMENT', 0)
    title.alignment = 1  # Center
    
    # Add main contract text
    for line in CONTRACT_TEMPLATE.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(10)
    
    # Add additional sections to reach the desired page count
    section_num = 10
    paragraphs_added = len(doc.paragraphs)
    target_paragraphs = num_pages * 35  # Approximate paragraphs per page
    
    while paragraphs_added < target_paragraphs:
        additional_text = ADDITIONAL_CLAUSE_TEMPLATE.format(section_num=section_num)
        for line in additional_text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(10)
                paragraphs_added += 1
        section_num += 1
        
        # Add some filler content if needed
        if paragraphs_added < target_paragraphs:
            filler = f"""
{section_num}. SUPPLEMENTARY PROVISIONS

{section_num}.1 This section contains additional terms and conditions that supplement the main agreement. 
The parties acknowledge that these provisions are essential to the proper execution and performance of 
this Agreement. Each party represents and warrants that it has carefully reviewed these provisions and 
understands their implications.

{section_num}.2 The parties further agree that any ambiguity in this Agreement shall be resolved in 
accordance with the principles of contract interpretation applicable in the governing jurisdiction. 
Neither party shall be deemed to be the drafter of this Agreement for purposes of interpreting any 
ambiguous provisions.
"""
            for line in filler.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.style.font.size = Pt(10)
                    paragraphs_added += 1
            section_num += 1
    
    doc.save(filename)
    print(f"Generated DOCX: {filename} (approx. {num_pages} pages)")


def generate_malformed_files():
    """Generate malformed files for error testing."""
    
    # Create a text file with .pdf extension (not a real PDF)
    with open('tests/fixtures/malformed_fake.pdf', 'w') as f:
        f.write("This is not a real PDF file, just plain text with a .pdf extension.")
    print("Generated malformed file: malformed_fake.pdf")
    
    # Create a text file with .docx extension (not a real DOCX)
    with open('tests/fixtures/malformed_fake.docx', 'w') as f:
        f.write("This is not a real DOCX file, just plain text with a .docx extension.")
    print("Generated malformed file: malformed_fake.docx")
    
    # Create an empty PDF file
    with open('tests/fixtures/empty.pdf', 'w') as f:
        f.write("")
    print("Generated malformed file: empty.pdf")
    
    # Create an empty DOCX file
    with open('tests/fixtures/empty.docx', 'w') as f:
        f.write("")
    print("Generated malformed file: empty.docx")
    
    # Create a file with unsupported extension
    with open('tests/fixtures/contract.txt', 'w') as f:
        f.write(CONTRACT_TEMPLATE)
    print("Generated unsupported file: contract.txt")
    
    # Create a corrupted PDF (partial PDF header)
    with open('tests/fixtures/corrupted.pdf', 'wb') as f:
        f.write(b'%PDF-1.4\n%\xE2\xE3\xCF\xD3\n')  # PDF header but incomplete
    print("Generated corrupted file: corrupted.pdf")


def main():
    """Generate all test contract files."""
    print("Generating test contract files...")
    print()
    
    # Create PDF contracts of various sizes
    generate_pdf_contract('tests/fixtures/contract_1page.pdf', 1)
    generate_pdf_contract('tests/fixtures/contract_10pages.pdf', 10)
    generate_pdf_contract('tests/fixtures/contract_25pages.pdf', 25)
    generate_pdf_contract('tests/fixtures/contract_50pages.pdf', 50)
    
    print()
    
    # Create DOCX contracts of various sizes
    generate_docx_contract('tests/fixtures/contract_1page.docx', 1)
    generate_docx_contract('tests/fixtures/contract_10pages.docx', 10)
    generate_docx_contract('tests/fixtures/contract_25pages.docx', 25)
    generate_docx_contract('tests/fixtures/contract_50pages.docx', 50)
    
    print()
    
    # Create malformed files
    generate_malformed_files()
    
    print()
    print("All test contract files generated successfully!")


if __name__ == '__main__':
    main()
