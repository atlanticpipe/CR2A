"""
Unit tests for ContractUploader class.
"""

import pytest
import tempfile
from pathlib import Path
from src.contract_uploader import ContractUploader
try:
    from pypdf import PdfWriter
except ImportError:
    from PyPDF2 import PdfWriter
from docx import Document


@pytest.fixture
def uploader():
    """Create a ContractUploader instance."""
    return ContractUploader()


@pytest.fixture
def temp_pdf_file():
    """Create a temporary PDF file with sample content."""
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.pdf', delete=False) as f:
        # Create a simple PDF with PyPDF2
        pdf_writer = PdfWriter()
        pdf_writer.add_blank_page(width=612, height=792)
        
        # Add a page with some text (note: blank pages don't have extractable text)
        # For testing, we'll create a minimal valid PDF
        pdf_writer.write(f)
        temp_path = f.name
    
    yield temp_path
    
    # Cleanup
    Path(temp_path).unlink(missing_ok=True)


@pytest.fixture
def temp_docx_file():
    """Create a temporary DOCX file with sample content."""
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    # Create a DOCX with python-docx
    doc = Document()
    doc.add_paragraph("Sample Contract")
    doc.add_paragraph("This is a test contract document.")
    doc.add_paragraph("Payment Terms: Net 30 days")
    doc.save(temp_path)
    
    yield temp_path
    
    # Cleanup
    Path(temp_path).unlink(missing_ok=True)


@pytest.fixture
def temp_txt_file():
    """Create a temporary TXT file (unsupported format)."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write("This is a text file")
        temp_path = f.name
    
    yield temp_path
    
    # Cleanup
    Path(temp_path).unlink(missing_ok=True)


@pytest.fixture
def temp_empty_file():
    """Create a temporary empty file."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.pdf', delete=False) as f:
        temp_path = f.name
    
    yield temp_path
    
    # Cleanup
    Path(temp_path).unlink(missing_ok=True)


class TestValidateFormat:
    """Tests for validate_format method."""
    
    def test_valid_pdf_format(self, uploader, temp_pdf_file):
        """Test that valid PDF files are accepted."""
        is_valid, error_msg = uploader.validate_format(temp_pdf_file)
        
        assert is_valid is True
        assert error_msg == ""
    
    def test_valid_docx_format(self, uploader, temp_docx_file):
        """Test that valid DOCX files are accepted."""
        is_valid, error_msg = uploader.validate_format(temp_docx_file)
        
        assert is_valid is True
        assert error_msg == ""
    
    def test_valid_txt_format(self, uploader, temp_txt_file):
        """Test that TXT files are now accepted."""
        is_valid, error_msg = uploader.validate_format(temp_txt_file)
        
        assert is_valid is True
        assert error_msg == ""
    
    def test_nonexistent_file(self, uploader):
        """Test that nonexistent files are rejected."""
        is_valid, error_msg = uploader.validate_format("nonexistent_file.pdf")
        
        assert is_valid is False
        assert "File not found" in error_msg
    
    def test_empty_file(self, uploader, temp_empty_file):
        """Test that empty files are rejected."""
        is_valid, error_msg = uploader.validate_format(temp_empty_file)
        
        assert is_valid is False
        assert "empty" in error_msg.lower()
    
    def test_directory_path(self, uploader, tmp_path):
        """Test that directory paths are rejected."""
        is_valid, error_msg = uploader.validate_format(str(tmp_path))
        
        assert is_valid is False
        assert "not a file" in error_msg.lower()


class TestGetFileInfo:
    """Tests for get_file_info method."""
    
    def test_pdf_file_info(self, uploader, temp_pdf_file):
        """Test extracting file info from PDF."""
        file_info = uploader.get_file_info(temp_pdf_file)
        
        assert 'filename' in file_info
        assert file_info['filename'].endswith('.pdf')
        assert 'file_size_bytes' in file_info
        assert file_info['file_size_bytes'] > 0
        assert 'file_size_mb' in file_info
        assert 'file_type' in file_info
        assert file_info['file_type'] == '.pdf'
        assert 'page_count' in file_info
    
    def test_docx_file_info(self, uploader, temp_docx_file):
        """Test extracting file info from DOCX."""
        file_info = uploader.get_file_info(temp_docx_file)
        
        assert 'filename' in file_info
        assert file_info['filename'].endswith('.docx')
        assert 'file_size_bytes' in file_info
        assert file_info['file_size_bytes'] > 0
        assert 'file_type' in file_info
        assert file_info['file_type'] == '.docx'
        assert 'page_count' in file_info
    
    def test_file_info_with_invalid_file(self, uploader):
        """Test get_file_info with nonexistent file returns minimal info."""
        file_info = uploader.get_file_info("nonexistent.pdf")
        
        # Should return minimal info without crashing
        assert 'filename' in file_info
        assert 'file_size_bytes' in file_info
        assert 'file_type' in file_info


class TestExtractText:
    """Tests for extract_text method."""
    
    def test_extract_text_from_docx(self, uploader, temp_docx_file):
        """Test text extraction from DOCX file."""
        text = uploader.extract_text(temp_docx_file)
        
        assert text is not None
        assert len(text) > 0
        assert "Sample Contract" in text
        assert "Payment Terms" in text
    
    def test_extract_text_from_txt_format(self, uploader, temp_txt_file):
        """Test that extracting text from TXT format works."""
        text = uploader.extract_text(temp_txt_file)
        
        assert text is not None
        assert len(text) > 0
        assert "This is a text file" in text
    
    def test_extract_text_from_nonexistent_file(self, uploader):
        """Test that extracting text from nonexistent file raises error."""
        with pytest.raises(ValueError) as exc_info:
            uploader.extract_text("nonexistent.pdf")
        
        assert "Cannot extract text" in str(exc_info.value)


class TestEdgeCases:
    """Tests for edge cases mentioned in task 3.5."""
    
    def test_empty_file_handling(self, uploader, temp_empty_file):
        """Test handling of empty files."""
        is_valid, error_msg = uploader.validate_format(temp_empty_file)
        
        assert is_valid is False
        assert "empty" in error_msg.lower()
    
    def test_large_file_size_check(self, uploader):
        """Test that file size limit is enforced."""
        # The MAX_FILE_SIZE is 100MB, we can verify the constant exists
        assert hasattr(uploader, 'MAX_FILE_SIZE')
        assert uploader.MAX_FILE_SIZE == 100 * 1024 * 1024


class TestMaxFileSizeParameter:
    """Tests for max_file_size parameter in __init__ method (Task 2.1)."""
    
    def test_default_max_file_size(self):
        """Test that default max file size is 100 MB for backward compatibility."""
        uploader = ContractUploader()
        
        assert uploader.MAX_FILE_SIZE == 100 * 1024 * 1024
    
    def test_custom_max_file_size(self):
        """Test that custom max file size is set correctly."""
        custom_size = 200 * 1024 * 1024  # 200 MB
        uploader = ContractUploader(max_file_size=custom_size)
        
        assert uploader.MAX_FILE_SIZE == custom_size
    
    def test_explicit_none_uses_default(self):
        """Test that explicitly passing None uses the default value."""
        uploader = ContractUploader(max_file_size=None)
        
        assert uploader.MAX_FILE_SIZE == 100 * 1024 * 1024
    
    def test_small_custom_size(self):
        """Test that small custom sizes work correctly."""
        small_size = 5 * 1024 * 1024  # 5 MB
        uploader = ContractUploader(max_file_size=small_size)
        
        assert uploader.MAX_FILE_SIZE == small_size
    
    def test_large_custom_size(self):
        """Test that large custom sizes work correctly."""
        large_size = 500 * 1024 * 1024  # 500 MB
        uploader = ContractUploader(max_file_size=large_size)
        
        assert uploader.MAX_FILE_SIZE == large_size
