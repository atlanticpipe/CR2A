from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any, Dict, Iterable, List, NamedTuple, Union

logger = logging.getLogger(__name__)

# Optional imports
try:
    # Only needed when backend="reportlab"
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
except Exception:
    colors = LETTER = ParagraphStyle = getSampleStyleSheet = inch = SimpleDocTemplate = Paragraph = Spacer = Table = TableStyle = None  # type: ignore

# ReportLab helpers
def _reportlab_styles():
    styles = getSampleStyleSheet()
    TitleStyle = ParagraphStyle("TitleStyle", parent=styles["Heading1"], spaceAfter=12)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=12, spaceAfter=6)
    H3 = ParagraphStyle("H3", parent=styles["Heading3"], spaceBefore=6, spaceAfter=4)
    Body = ParagraphStyle("Body", parent=styles["BodyText"], leading=14, spaceAfter=6)
    Small = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=9, leading=12)
    return TitleStyle, H2, H3, Body, Small

def _kv_table_reportlab(section_i: Dict[str, Any], Body) -> "Table":
    rows: List[List[Any]] = []
    for k, v in section_i.items():
        key = k[:-1] if k.endswith(":") else k
        rows.append([Paragraph(f"<b>{key}</b>", Body), Paragraph(str(v or ""), Body)])
    t = Table(rows, colWidths=[2.2 * inch, 4.8 * inch])
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
            ]
        )
    )
    return t

def _add_clause_block_reportlab(story: List[Any], block: Dict[str, Any], H3, Body) -> None:
    # Accept both schema-normalized keys (with spaces) and analyzer keys (snake_case) so PDFs render data instead of blanks.
    fields = [
        ("Clause Language", ["Clause Language", "clause_language"]),
        ("Clause Summary", ["Clause Summary", "clause_summary"]),
        ("Risk Triggers", ["Risk Triggers Identified", "risk_triggers"]),
        ("Flow-Down Obligations", ["Flow-Down Obligations", "flow_down_obligations"]),
        ("Redline Recommendations", ["Redline Recommendations", "redline_recommendations"]),
        ("Harmful Language / Conflicts", ["Harmful Language / Policy Conflicts", "harmful_language_conflicts"]),
    ]

    def _first_present(obj: Dict[str, Any], keys: List[str]) -> Any:
        # Pick the first populated value to avoid empty paragraphs when the other shape is absent.
        for k in keys:
            val = obj.get(k)
            if val:
                return val
        return None

    for label, keys in fields:
        val = _first_present(block, keys)
        if val:
            story.append(Paragraph(f"<b>{label}</b>", H3))
            story.append(Paragraph(str(val), Body))

def _section_items_reportlab(story: List[Any], section_name: str, items: Iterable[Dict[str, Any]], H2, H3, Body, Small) -> None:
    story.append(Paragraph(section_name, H2))
    empty = True
    for item in items or []:
        empty = False
        title = item.get("item_title") or f"{section_name} Item"
        story.append(Paragraph(title, H3))
        for block in item.get("clauses") or []:
            _add_clause_block_reportlab(story, block, H3, Body)
        story.append(Spacer(1, 6))
    if empty:
        story.append(Paragraph("No items found.", Small))

def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.drawRightString(LETTER[0] - 0.5 * inch, 0.5 * inch, f"Page {doc.page}")
    canvas.restoreState()

def _export_reportlab(data: Dict[str, Any], output_pdf: Path, title: str) -> Path:
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab is not installed; install it or use backend='docx'")
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    TitleStyle, H2, H3, Body, Small = _reportlab_styles()

    doc = SimpleDocTemplate(
        str(output_pdf),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title=title,
        author="CR2A",
    )

    story: List[Any] = []
    story.append(Paragraph(title, TitleStyle))
    story.append(Spacer(1, 6))

    section_i = data.get("SECTION_I") or {}
    story.append(Paragraph("SECTION I — Contract Overview", H2))
    story.append(_kv_table_reportlab(section_i, Body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("SECTION II — Parties & Dates", H2))
    _section_items_reportlab(story, "Parties & Dates", data.get("SECTION_II") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION III — Scope & Deliverables", H2))
    _section_items_reportlab(story, "Scope & Deliverables", data.get("SECTION_III") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION IV — Payment Terms", H2))
    _section_items_reportlab(story, "Payment Terms", data.get("SECTION_IV") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION V — Risk Allocation", H2))
    _section_items_reportlab(story, "Risk Allocation", data.get("SECTION_V") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    story.append(Paragraph("SECTION VI — Compliance & Legal", H2))
    _section_items_reportlab(story, "Compliance & Legal", data.get("SECTION_VI") or [], H2, H3, Body, Small)
    story.append(Spacer(1, 6))

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return output_pdf


def _resolve_fillable_template(map_path: Path, template_override: Union[str, Path, None]) -> Path:
    # Resolve the fillable PDF template path, preferring an explicit override then the map hint.
    map_dir = map_path.parent
    template_hint = Path(template_override) if template_override else None
    if template_hint is None:
        data = json.loads(map_path.read_text(encoding="utf-8"))
        template_hint = Path(data.get("template_path", ""))
    if not template_hint:
        raise FileNotFoundError("No fillable PDF template path provided in pdf_field_map.json or function args")
    candidate = template_hint if template_hint.is_absolute() else (map_dir / template_hint)
    if not candidate.exists():
        raise FileNotFoundError(f"Fillable PDF template not found at {candidate}")
    return candidate


def _build_fillable_fields(data: Dict[str, Any], field_map: Dict[str, Any]) -> Dict[str, str]:
    # Translate normalized JSON into flat PDF form field/value pairs using the provided mapping.
    values: Dict[str, str] = {}

    section_i_map = field_map.get("SECTION_I", {})
    section_i_data = data.get("SECTION_I") or {}
    for json_key, field_name in section_i_map.items():
        if field_name:
            values[field_name] = str(section_i_data.get(json_key, ""))

    def _fill_section(section_key: str) -> None:
        section_map = field_map.get(section_key) or {}
        items = data.get(section_key) or []
        for item_idx, item in enumerate(items, start=1):
            title_field = section_map.get("item_title")
            if title_field:
                values[title_field.format(item_index=item_idx, section=section_key)] = str(item.get("item_title", ""))
            closing_field = section_map.get("closing_line")
            if closing_field and item.get("closing_line"):
                values[closing_field.format(item_index=item_idx, section=section_key)] = str(item.get("closing_line"))

            clause_map = section_map.get("clauses") or {}
            for clause_idx, clause in enumerate(item.get("clauses") or [], start=1):
                for label, field_template in clause_map.items():
                    val = clause.get(label)
                    if val is None:
                        # Permit snake_case analyzer keys in case normalization is bypassed.
                        val = clause.get(label.replace(" ", "_").lower())
                    if val:
                        field_name = field_template.format(
                            section=section_key,
                            item_index=item_idx,
                            clause_index=clause_idx,
                        )
                        values[field_name] = str(val)

    for sec in ["SECTION_II", "SECTION_III", "SECTION_IV", "SECTION_V", "SECTION_VI"]:
        _fill_section(sec)

    return values


def _export_fillable_pdf(
    data: Dict[str, Any],
    output_pdf: Path,
    *,
    field_map_path: Union[str, Path, None] = None,
    template_pdf: Union[str, Path, None] = None,
) -> Path:
    # Populate a fillable PDF form using pypdf; falls back to caller for alternate rendering on failure.
    try:
        from pypdf import PdfReader, PdfWriter
        from pypdf.generic import BooleanObject, NameObject
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("pypdf is required for backend='fillable_pdf'") from exc

    map_path = Path(field_map_path).expanduser().resolve() if field_map_path else Path(__file__).resolve().parents[2] / "templates" / "pdf_field_map.json"
    if not map_path.exists():
        raise FileNotFoundError(f"Field map not found: {map_path}")

    try:
        field_map = json.loads(map_path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise ValueError(f"Invalid pdf_field_map.json: {exc}") from exc

    template_path = _resolve_fillable_template(map_path, template_pdf)

    field_values = _build_fillable_fields(data, field_map.get("fields", {}))
    if not field_values:
        raise ValueError("No field mappings resolved; pdf_field_map.json produced an empty payload")

    reader = PdfReader(str(template_path))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)

    for page in writer.pages:
        writer.update_page_form_field_values(page, field_values)

    # Hint to PDF viewers to render updated appearances so filled values are visible.
    if "/AcroForm" in writer._root_object:
        writer._root_object[NameObject("/AcroForm")][NameObject("/NeedAppearances")] = BooleanObject(True)

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    with output_pdf.open("wb") as fh:
        writer.write(fh)
    return output_pdf

# DOCX backend
def _export_docx(data: Dict[str, Any], template_docx: Path, output_pdf: Path, title: str) -> Path:
    from docx import Document
    try:
        from docx2pdf import convert as docx2pdf_convert
    except Exception:
        docx2pdf_convert = None  # optional

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    if template_docx and template_docx.exists():
        doc = Document(str(template_docx))
    else:
        doc = Document()
        doc.add_heading(title, level=1)

    # SECTION I
    doc.add_heading("SECTION I — Contract Overview", level=2)
    section_i = data.get("SECTION_I") or {}
    table = doc.add_table(rows=0, cols=2)
    for k, v in section_i.items():
        row = table.add_row().cells
        row[0].text = (k[:-1] if k.endswith(":") else k)
        row[1].text = str(v or "")

    def write_items(hdr: str, items: Iterable[Dict[str, Any]]):
        doc.add_heading(hdr, level=2)
        wrote = False
        for item in items or []:
            wrote = True
            doc.add_heading(str(item.get("item_title") or "Item"), level=3)
            for block in item.get("clauses") or []:
                for label, key in [
                    ("Clause Language", "clause_language"),
                    ("Clause Summary", "clause_summary"),
                    ("Risk Triggers", "risk_triggers"),
                    ("Flow-Down Obligations", "flow_down_obligations"),
                    ("Redline Recommendations", "redline_recommendations"),
                    ("Harmful Language / Conflicts", "harmful_language_conflicts"),
                ]:
                    val = block.get(key)
                    if val:
                        p = doc.add_paragraph()
                        run = p.add_run(f"{label}: ")
                        run.bold = True
                        doc.add_paragraph(str(val))
        if not wrote:
            doc.add_paragraph("No items found.")

    write_items("SECTION II — Parties & Dates", data.get("SECTION_II"))
    write_items("SECTION III — Scope & Deliverables", data.get("SECTION_III"))
    write_items("SECTION IV — Payment Terms", data.get("SECTION_IV"))
    write_items("SECTION V — Risk Allocation", data.get("SECTION_V"))
    write_items("SECTION VI — Compliance & Legal", data.get("SECTION_VI"))

    out_docx = output_pdf.with_suffix(".docx")
    doc.save(str(out_docx))

    # Convert to PDF if possible; otherwise leave the .docx
    try:
        from docx2pdf import convert as docx2pdf_convert2
        if docx2pdf_convert2:
            docx2pdf_convert2(str(out_docx), str(output_pdf))
            return output_pdf
    except Exception:
        pass
    return out_docx

# Public API
def export_pdf_from_filled_json(
    data: Dict[str, Any],
    output_pdf: Union[str, Path],
    *,
    backend: str = "reportlab",  # default to ReportLab to avoid OS-level conversion
    template_docx: Union[str, Path, None] = None,
    title: str = "Contract Risk & Compliance Analysis",
) -> Path:
    """
    Render a CR2A filled JSON into a PDF (or DOCX if conversion is unavailable).
    backend="docx" uses python-docx (+ docx2pdf if available); backend="reportlab" uses ReportLab.
    """
    return render_report(
        data,
        output_pdf,
        backend=backend,
        template_docx=template_docx,
        title=title,
    ).path


class RenderResult(NamedTuple):
    path: Path
    backend: str
    fallback_reason: str | None


def render_report(
    data: Dict[str, Any],
    output_pdf: Union[str, Path],
    *,
    backend: str = "reportlab",
    template_docx: Union[str, Path, None] = None,
    template_pdf: Union[str, Path, None] = None,
    field_map_path: Union[str, Path, None] = None,
    title: str = "Contract Risk & Compliance Analysis",
) -> RenderResult:
    """
    Render normalized CR2A JSON with a pluggable backend.
    - backend="reportlab" (default): paginated tables using ReportLab
    - backend="docx": template-based DOCX converted to PDF when possible
    - backend="fillable_pdf": populate a fillable PDF form using a mapping file
    Gracefully falls back to ReportLab if the requested backend is unavailable.
    """
    output_pdf = Path(output_pdf).expanduser().resolve()
    backend_lc = (backend or "reportlab").lower()

    if backend_lc == "docx":
        tpl = Path(template_docx) if template_docx else Path("")
        return RenderResult(_export_docx(data, tpl, output_pdf, title), "docx", None)

    if backend_lc == "fillable_pdf":
        try:
            return RenderResult(
                _export_fillable_pdf(data, output_pdf, field_map_path=field_map_path, template_pdf=template_pdf),
                "fillable_pdf",
                None,
            )
        except Exception as exc:
            logger.warning(
                "fillable_pdf backend unavailable (%s); falling back to reportlab for %s",
                exc,
                output_pdf,
            )
            return RenderResult(_export_reportlab(data, output_pdf, title), "reportlab", str(exc))

    # Default path is ReportLab to minimize external dependencies.
    if backend_lc != "reportlab":
        raise ValueError(f"Unsupported backend: {backend}")
    return RenderResult(_export_reportlab(data, output_pdf, title), "reportlab", None)